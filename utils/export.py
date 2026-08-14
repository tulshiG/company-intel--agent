"""
utils/export.py
------------------
Turns an IntelligenceReport into downloadable artifacts. Markdown is the
canonical representation; PDF and DOCX are both generated FROM the
markdown-equivalent structure so formatting stays consistent across
formats.
"""
from io import BytesIO
from agent.models import IntelligenceReport


def to_markdown(report: IntelligenceReport) -> str:
    r = report
    lines = [
        f"# Company Intelligence Report: {r.company_name}",
        f"*Generated {r.generated_at} UTC · LLM: {r.llm_provider} · Search: {r.search_backend}*",
        "",
        "## 1. Company Overview",
        r.overview.summary,
        "",
        f"- **Industry:** {r.overview.industry}",
        f"- **Scale:** {r.overview.scale}",
        f"- **Geographic Presence:** {r.overview.geographic_presence}",
        "",
        "## 2. Key Business Information",
        "**Major Offerings**",
        *[f"- {x}" for x in r.key_info.major_offerings],
        "",
        "**Recent Developments**",
        *[f"- {x}" for x in r.key_info.recent_developments],
        "",
        "**Expansion Plans**",
        *[f"- {x}" for x in r.key_info.expansion_plans],
        "",
        "**Other Public Info**",
        *[f"- {x}" for x in r.key_info.public_info_highlights],
        "",
        "## 3. Potential Business Challenges",
    ]
    for c in r.challenges:
        lines += [
            f"### {c.title}  *({c.category})*",
            c.description,
            f"> **Reasoning:** {c.reasoning}",
            "",
        ]
    lines.append("## 4. AI Opportunities")
    for o in r.ai_opportunities:
        linked = f"  \n*Addresses: {o.linked_challenge}*" if o.linked_challenge else ""
        lines += [
            f"### {o.title}  *({o.area})*",
            o.description,
            f"**Expected impact:** {o.expected_impact}{linked}",
            "",
        ]
    lines += [
        "## 5. Personalized Pitch",
        r.pitch,
        "",
        "## Sources",
        *[f"- [{s.title}]({s.url})" for s in r.sources],
    ]
    return "\n".join(lines)


def to_pdf_bytes(report: IntelligenceReport) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def cell(text, h=6, new_x=XPos.LMARGIN, new_y=YPos.NEXT):
        # multi_cell with w=0 otherwise leaves the cursor at the right
        # margin instead of resetting to the left margin on the next line.
        pdf.multi_cell(0, h, _clean(text), new_x=new_x, new_y=new_y)

    pdf.set_font("Helvetica", "B", 16)
    cell(f"Company Intelligence Report: {report.company_name}", h=10)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    cell(f"Generated {report.generated_at} UTC | LLM: {report.llm_provider} | Search: {report.search_backend}")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(2)

    def h2(text):
        pdf.set_font("Helvetica", "B", 13)
        pdf.ln(4)
        cell(text, h=8)
        pdf.set_font("Helvetica", "", 10)

    def h3(text):
        pdf.set_font("Helvetica", "B", 11)
        pdf.ln(2)
        cell(text, h=7)
        pdf.set_font("Helvetica", "", 10)

    def body(text):
        cell(text)

    def bullet(text):
        cell(f"-  {text}")

    h2("1. Company Overview")
    body(report.overview.summary)
    bullet(f"Industry: {report.overview.industry}")
    bullet(f"Scale: {report.overview.scale}")
    bullet(f"Geographic Presence: {report.overview.geographic_presence}")

    h2("2. Key Business Information")
    h3("Major Offerings")
    for x in report.key_info.major_offerings:
        bullet(x)
    h3("Recent Developments")
    for x in report.key_info.recent_developments:
        bullet(x)
    h3("Expansion Plans")
    for x in report.key_info.expansion_plans:
        bullet(x)
    h3("Other Public Info")
    for x in report.key_info.public_info_highlights:
        bullet(x)

    h2("3. Potential Business Challenges")
    for c in report.challenges:
        h3(f"{c.title} ({c.category})")
        body(c.description)
        body(f"Reasoning: {c.reasoning}")

    h2("4. AI Opportunities")
    for o in report.ai_opportunities:
        h3(f"{o.title} ({o.area})")
        body(o.description)
        body(f"Expected impact: {o.expected_impact}")
        if o.linked_challenge:
            body(f"Addresses: {o.linked_challenge}")

    h2("5. Personalized Pitch")
    body(report.pitch)

    if report.sources:
        h2("Sources")
        for s in report.sources:
            bullet(f"{s.title} - {s.url}")

    out = pdf.output()
    return bytes(out)


def to_docx_bytes(report: IntelligenceReport) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(f"Company Intelligence Report: {report.company_name}", level=0)
    doc.add_paragraph(
        f"Generated {report.generated_at} UTC | LLM: {report.llm_provider} | Search: {report.search_backend}"
    ).italic = True

    doc.add_heading("1. Company Overview", level=1)
    doc.add_paragraph(report.overview.summary)
    doc.add_paragraph(f"Industry: {report.overview.industry}", style="List Bullet")
    doc.add_paragraph(f"Scale: {report.overview.scale}", style="List Bullet")
    doc.add_paragraph(f"Geographic Presence: {report.overview.geographic_presence}", style="List Bullet")

    doc.add_heading("2. Key Business Information", level=1)
    for label, items in [
        ("Major Offerings", report.key_info.major_offerings),
        ("Recent Developments", report.key_info.recent_developments),
        ("Expansion Plans", report.key_info.expansion_plans),
        ("Other Public Info", report.key_info.public_info_highlights),
    ]:
        doc.add_heading(label, level=2)
        for x in items:
            doc.add_paragraph(x, style="List Bullet")

    doc.add_heading("3. Potential Business Challenges", level=1)
    for c in report.challenges:
        doc.add_heading(f"{c.title} ({c.category})", level=2)
        doc.add_paragraph(c.description)
        p = doc.add_paragraph()
        p.add_run("Reasoning: ").bold = True
        p.add_run(c.reasoning)

    doc.add_heading("4. AI Opportunities", level=1)
    for o in report.ai_opportunities:
        doc.add_heading(f"{o.title} ({o.area})", level=2)
        doc.add_paragraph(o.description)
        p = doc.add_paragraph()
        p.add_run("Expected impact: ").bold = True
        p.add_run(o.expected_impact)
        if o.linked_challenge:
            doc.add_paragraph(f"Addresses: {o.linked_challenge}")

    doc.add_heading("5. Personalized Pitch", level=1)
    for para in report.pitch.split("\n"):
        if para.strip():
            doc.add_paragraph(para)

    if report.sources:
        doc.add_heading("Sources", level=1)
        for s in report.sources:
            doc.add_paragraph(f"{s.title} - {s.url}", style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _clean(text: str) -> str:
    """fpdf2's Helvetica core font is latin-1 only; strip characters it can't render."""
    if not text:
        return ""
    return text.encode("latin-1", errors="replace").decode("latin-1")
